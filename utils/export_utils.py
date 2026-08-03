"""
Export pipeline results (title, summary, action items, decisions, questions,
transcript) to downloadable PDF or DOCX files, in-memory (no disk writes needed).
"""
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF


def _sections(result: dict):
    """Common ordered list of (heading, content) pairs used by both exporters."""
    return [
        ("Title", result.get("title", "")),
        ("Summary", result.get("summary", "")),
        ("Action Items", result.get("action_items", "")),
        ("Key Decisions", result.get("key_decisions", "")),
        ("Open Questions", result.get("open_questions", "")),
        ("Full Transcript", result.get("transcript", "")),
    ]


def export_to_docx(result: dict) -> bytes:
    """Build a .docx file in memory and return its raw bytes."""
    doc = Document()

    title = doc.add_heading("AI Video Assistant — Meeting Report", level=0)
    title.runs[0].font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True

    for heading, content in _sections(result):
        doc.add_heading(heading, level=1)
        para = doc.add_paragraph(content or "—")
        para.style.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_to_pdf(result: dict) -> bytes:
    """Build a .pdf file in memory and return its raw bytes."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(15, 15, 15)
    pdf.add_page()

    def write_block(text, font="Helvetica", style="", size=10, color=(50, 50, 60), gap=0):
        pdf.set_x(pdf.l_margin)
        pdf.set_font(font, style, size)
        pdf.set_text_color(*color)
        pdf.multi_cell(0, 6, text, new_x="LMARGIN", new_y="NEXT")
        if gap:
            pdf.ln(gap)

    write_block("AI Video Assistant - Meeting Report", style="B", size=18, color=(124, 58, 237), gap=2)
    write_block(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", style="I", size=9, color=(100, 100, 100), gap=4)

    for heading, content in _sections(result):
        write_block(heading, style="B", size=13, color=(20, 20, 30))
        safe_text = (content or "-").encode("latin-1", "replace").decode("latin-1")
        write_block(safe_text, size=10, color=(50, 50, 60), gap=3)

    # fpdf2's output() returns a bytearray; wrap for Streamlit's download_button
    return bytes(pdf.output())
